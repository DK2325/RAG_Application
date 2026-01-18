import streamlit as st
import requests
import numpy as np
import faiss
from PyPDF2 import PdfReader
from docx import Document
from io import StringIO
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
import os
import re

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

EMBED_MODEL = os.getenv("EMBED_MODEL", "nvidia/llama-3.2-nv-embedqa-1b-v2")
LLM_MODEL = os.getenv("LLM_MODEL", "meta/llama-3.3-70b-instruct")

# ---------------- FILE READERS ---------------- #

def read_pdf(file):
    reader = PdfReader(file)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def read_txt(file):
    return StringIO(file.getvalue().decode("utf-8")).read()

def read_docx(file):
    doc = Document(file)
    return "\n".join([p.text for p in doc.paragraphs])

def read_data_txt():
    return "RAG stands for Retrieval-Augmented Generation..."

# ---------------- TASK DETECTION ---------------- #

def is_task_question(question: str) -> bool:
    q = question.lower().strip()
    task_verbs = [
        "summarize", "summary", "explain", "rewrite", "rephrase",
        "extract", "list", "analyze", "analyse", "compare",
        "format", "convert", "simplify", "improve",
        "generate", "create", "outline", "describe"
    ]
    return any(v in q for v in task_verbs)

# ---------------- EMBEDDINGS ---------------- #

def embed_chunks(chunks):
    url = "https://integrate.api.nvidia.com/v1/embeddings"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {os.getenv('NVIDIA_API_KEY')}"
    }
    payload = {"input": chunks, "model": EMBED_MODEL, "input_type": "passage"}
    response = requests.post(url, json=payload, headers=headers, timeout=30)
    response.raise_for_status()
    return [item["embedding"] for item in response.json()["data"]]

def get_query_embedding(question):
    url = "https://integrate.api.nvidia.com/v1/embeddings"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {os.getenv('NVIDIA_API_KEY')}"
    }
    payload = {"input": [question], "model": EMBED_MODEL, "input_type": "query"}
    response = requests.post(url, json=payload, headers=headers, timeout=20)
    response.raise_for_status()
    return response.json()["data"][0]["embedding"]

# ---------------- LLM CALLS ---------------- #

def run_task_on_document(document_text, task):
    url = "https://integrate.api.nvidia.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {os.getenv('NVIDIA_API_KEY')}"
    }

    prompt = f"""
You are an AI assistant working ONLY with the given document.

Perform the user's task using the document content.
Do NOT say that information is missing.
Do NOT refuse.
Always attempt the task using what is available.

Document:
{document_text}

Task:
{task}
"""

    payload = {
        "model": LLM_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 2048
    }

    response = requests.post(url, headers=headers, json=payload, timeout=60)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def ask_llm_fact(context, question):
    url = "https://integrate.api.nvidia.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {os.getenv('NVIDIA_API_KEY')}"
    }

    prompt = f"""
Answer the question using ONLY the context below.
If the answer truly does not exist, say "Not mentioned in the document".

Context:
{context}

Question:
{question}
"""

    payload = {
        "model": LLM_MODEL,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": 2048
    }

    response = requests.post(url, headers=headers, json=payload, timeout=60)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

# ---------------- STREAMLIT UI ---------------- #

st.title("Task-Aware RAG System")

uploaded_file = st.file_uploader(
    "Upload PDF / TXT / DOCX (optional)",
    type=["pdf", "txt", "docx"]
)

question = st.text_input("Ask a question or give a task:")
use_doc_only = st.checkbox("Use only the uploaded document", value=bool(uploaded_file))

if st.button("Generate Answer") and question.strip():

    with st.spinner("Reading document..."):
        if uploaded_file:
            if uploaded_file.type == "application/pdf":
                document_text = read_pdf(uploaded_file)
            elif uploaded_file.type == "text/plain":
                document_text = read_txt(uploaded_file)
            else:
                document_text = read_docx(uploaded_file)
        else:
            document_text = read_data_txt()

    # 🔥 TASK MODE (NO RETRIEVAL)
    if is_task_question(question):
        with st.spinner("Performing task on document..."):
            result = run_task_on_document(document_text, question)
            st.success("Result")
            st.write(result)
        st.stop()

    # -------- FACT MODE (RAG) -------- #

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    docs = splitter.split_documents([LCDocument(page_content=document_text)])
    texts = [d.page_content for d in docs]

    with st.spinner("Embedding document..."):
        embeddings = embed_chunks(texts)
        index = faiss.IndexFlatL2(len(embeddings[0]))
        index.add(np.array(embeddings).astype("float32"))

    with st.spinner("Searching document..."):
        q_emb = get_query_embedding(question)
        D, I = index.search(np.array([q_emb]).astype("float32"), k=5)
        context = "\n\n---\n\n".join(texts[i] for i in I[0])

    with st.spinner("Generating answer..."):
        answer = ask_llm_fact(context, question)
        st.success("Answer")
        st.write(answer)
